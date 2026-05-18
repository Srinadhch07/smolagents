import requests
import os
from smolagents import tool, CodeAgent, InferenceClientModel, LiteLLMModel
from dotenv import load_dotenv
from urllib.parse import quote
from docx import Document   
from reportlab.platypus import ( SimpleDocTemplate, Paragraph, Spacer)
from reportlab.lib.styles import getSampleStyleSheet
import litellm

load_dotenv()
litellm._turn_on_debug()

hf_token = os.getenv("HF_TOKEN")
groq_token = os.getenv("GROQ_API_KEY")
USE_FAST_MODEL = os.getenv("USE_FAST_MODE", True)

if not hf_token:
    raise ValueError("please set the env varibale")
if USE_FAST_MODEL:
    model = LiteLLMModel( model_id="groq/llama-3.3-70b-versatile", api_key=groq_token)
else:
    model = InferenceClientModel( model_id="Qwen/Qwen2.5-Coder-32B-Instruct", token=hf_token)


@tool
def get_weather(city: str) -> str:
    """
    Returns the current weather forecast for a specified city.
    Args:
        city: The name of the city to get the weather for.
    """
    city_encoded = quote(city)
    response = requests.get(f"https://wttr.in/{city_encoded}?format=%C+%t", timeout=10)
    if response.status_code == 200:
        return f"The weather in {city} is: {response.text.strip()}"
    else:
        return "Sorry, I couldn't fetch the weather data."
    
@tool
def save_to_file(content: str, filename: str = "weather_report.txt") -> str:
    """
    Saves the provided text content to a file.
    Args:
        content: The text content to save.
        filename: The name of the file to save to (default: weather_report.txt).
    """
    with open(filename, "w") as f:
        f.write(content)
    return f"Content successfully saved to {filename}"


@tool
def search_jobs( keyword: str, location: str = "India" ) -> str:
    """
    Searches jobs using the Adzuna API.

    Args:
        keyword: Job title or skill to search for.
        location: Job location like Bangalore, Hyderabad, Remote, etc.
    """
    app_id = os.getenv("ADZUNA_APP_ID")
    app_key = os.getenv("ADZUNA_APP_KEY")
    if not app_id or not app_key:
        return "Missing Adzuna API credentials."

    url = "https://api.adzuna.com/v1/api/jobs/in/search/1"

    params = {
        "app_id": app_id,
        "app_key": app_key,
        "what": keyword,
        "where": location,
        "results_per_page": 5,
        "content-type": "application/json"
    }

    response = requests.get(
        url,
        params=params,
        timeout=10
    )

    if response.status_code != 200:
        return f"API Error: {response.status_code}"

    data = response.json()

    jobs = data.get("results", [])

    if not jobs:
        return "No jobs found."

    formatted_jobs = []

    for idx, job in enumerate(jobs, start=1):

        title = job.get("title", "N/A")

        company = job.get(
            "company",
            {}
        ).get(
            "display_name",
            "Unknown"
        )

        location = job.get(
            "location",
            {}
        ).get(
            "display_name",
            "Unknown"
        )

        salary_min = job.get("salary_min")
        salary_max = job.get("salary_max")

        redirect_url = job.get("redirect_url", "N/A")

        salary_text = "Not specified"

        if salary_min and salary_max:
            salary_text = (
                f"{int(salary_min):,} - "
                f"{int(salary_max):,}"
            )

        formatted_jobs.append(
            f"""
Job {idx}
Title: {title}
Company: {company}
Location: {location}
Salary: {salary_text}
Apply: {redirect_url}
"""
        )

    return "\n".join(formatted_jobs)

@tool
def save_markdown(
    content: str,
    filename: str = "report.md"
) -> str:
    """
    Saves content as a markdown file.

    Args:
        content: Markdown content to save.
        filename: Markdown filename.
    """

    with open(filename, "w", encoding="utf-8") as f:
        f.write(content)

    return f"Markdown saved to {filename}"
    
@tool
def save_docx(
    content: str,
    filename: str = "report.docx"
) -> str:
    """
    Saves content as a DOCX file.

    Args:
        content: Text content.
        filename: DOCX filename.
    """

    doc = Document()

    doc.add_heading("Generated Report", level=1)

    doc.add_paragraph(content)

    doc.save(filename)

    return f"DOCX saved to {filename}"


@tool
def save_pdf(
    content: str,
    filename: str = "report.pdf"
) -> str:
    """
    Saves content as a PDF file.

    Args:
        content: Text content.
        filename: PDF filename.
    """

    doc = SimpleDocTemplate(filename)

    styles = getSampleStyleSheet()

    elements = []

    elements.append(
        Paragraph(content, styles["BodyText"])
    )

    elements.append(Spacer(1, 12))

    doc.build(elements)

    return f"PDF saved to {filename}"



agent = CodeAgent(
    tools=[get_weather, 
        save_to_file, 
        search_jobs,
        save_markdown,
        save_docx,
        save_pdf],
    model=model,
    add_base_tools=True,
    verbosity_level=2
)


if __name__ == "__main__":
    prompt = input("You: ")
    response = agent.run(prompt)
    print(f'Agent: {response}')